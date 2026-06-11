"""
Document uploader — handles file uploads to Supabase Storage and
text extraction from Markdown and plain-text files.
"""

import logging
from fastapi import HTTPException
from supabase_db import get_client

logger = logging.getLogger(__name__)


def _ensure_bucket_exists(client, bucket: str) -> None:
    """Crea el bucket si no existe. Ignora el error si ya existe."""
    try:
        client.storage.get_bucket(bucket)
    except Exception as e:
        err_str = str(e).lower()
        # Si el error indica que no existe, intentar crearlo
        if "not found" in err_str or "does not exist" in err_str or "404" in err_str:
            try:
                client.storage.create_bucket(bucket, options={"public": False})
                logger.info(f"Bucket '{bucket}' creado exitosamente.")
            except Exception as create_err:
                create_err_str = str(create_err).lower()
                # Si ya existe (race condition), ignorar
                if "already exists" in create_err_str or "duplicate" in create_err_str or "409" in create_err_str:
                    pass
                else:
                    raise HTTPException(
                        status_code=500,
                        detail=f"No se pudo crear el bucket de almacenamiento: {str(create_err)}"
                    )
        else:
            # Otro error al verificar el bucket (ej: permisos, red)
            logger.warning(f"Advertencia al verificar bucket '{bucket}': {e}. Intentando upload de todas formas.")


async def upload_file_to_blob(content: bytes, blob_path: str, content_type: str) -> str:
    client = get_client()
    bucket = "documents"

    # Garantizar que el bucket existe antes de subir
    _ensure_bucket_exists(client, bucket)

    # Subir archivo con manejo de errores explícito
    try:
        client.storage.from_(bucket).upload(
            path=blob_path,
            file=content,
            file_options={"content-type": content_type, "upsert": "true"},
        )
    except Exception as e:
        err_str = str(e).lower()
        logger.error(f"Error al subir archivo a Supabase Storage (path={blob_path}): {e}")
        if "duplicate" in err_str or "already exists" in err_str or "409" in err_str:
            # El archivo ya existe y upsert no funcionó — no es un error crítico
            logger.warning(f"Archivo ya existe en storage, continuando: {blob_path}")
        else:
            raise HTTPException(
                status_code=500,
                detail=f"Error al guardar el archivo en el almacenamiento: {str(e)}"
            )

    # Construir URL de referencia (el bucket es privado, así que es una URL de acceso interno)
    return client.storage.from_(bucket).get_public_url(blob_path)


def download_from_blob(blob_path: str) -> bytes:
    client = get_client()
    bucket = "documents"
    # Extract path inside bucket
    return client.storage.from_(bucket).download(blob_path)


def extract_text_from_file(content: bytes, filename: str, content_type: str | None) -> str:
    """
    Extract plain text from uploaded files.
    Supports: Markdown (.md), plain text (.txt), PDF (.pdf), and Word (.docx).
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = []
            for page in doc:
                t = page.get_text()
                if t:
                    text.append(t)
            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Error al extraer texto del PDF: {str(e)}")

    elif lower_name.endswith(".docx"):
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            text = []
            
            # Extraer texto de párrafos
            for p in doc.paragraphs:
                if p.text:
                    text.append(p.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    seen_cells = set()
                    for cell in row.cells:
                        if cell not in seen_cells:
                            seen_cells.add(cell)
                            val = cell.text.strip()
                            if val:
                                row_text.append(val)
                    if row_text:
                        text.append(" | ".join(row_text))
                        
            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Error al extraer texto del archivo DOCX: {str(e)}")

    elif lower_name.endswith((".md", ".txt")) or (content_type and content_type in ("text/markdown", "text/plain")):
        return content.decode("utf-8", errors="replace")

    # Fallback: try as UTF-8 text anyway
    return content.decode("utf-8", errors="replace")
